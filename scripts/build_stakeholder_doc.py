"""Regenerate deliverables/stakeholder_status_update.docx from the markdown.

Markdown is the source of truth. The .docx is a derived artifact (gitignored).

Usage:
    python scripts/build_stakeholder_doc.py
    # or:
    make stakeholder-doc

Tries pandoc first (better quality output for tables/headings); falls back to
python-docx if pandoc isn't on PATH (less faithful but works without
external dependencies).
"""

from __future__ import annotations

import os
import shutil
import subprocess
import sys

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_MD = os.path.join(REPO, "deliverables", "stakeholder_status_update.md")
OUT_DOCX = os.path.join(REPO, "deliverables", "stakeholder_status_update.docx")


def run_pandoc() -> bool:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False
    cmd = [pandoc, SRC_MD, "--from", "gfm", "--to", "docx", "-o", OUT_DOCX]
    print(f"$ {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"pandoc failed (exit {result.returncode}):\n{result.stderr}",
              file=sys.stderr)
        return False
    return True


def run_python_docx() -> bool:
    try:
        import docx  # noqa: F401
    except ImportError:
        print("Installing python-docx (one-time)...")
        subprocess.check_call([sys.executable, "-m", "pip", "install",
                               "--quiet", "python-docx"])
    # Import after install
    from docx import Document

    doc = Document()
    with open(SRC_MD, encoding="utf-8") as f:
        text = f.read()

    # Very crude markdown-to-docx (paragraph and heading level only). Use
    # pandoc for tables/lists fidelity. This fallback is just a backstop so
    # the build always produces SOMETHING.
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == "---":
            doc.add_paragraph("")  # rule placeholder
        else:
            doc.add_paragraph(line)
    doc.save(OUT_DOCX)
    return True


def main() -> int:
    if not os.path.exists(SRC_MD):
        print(f"ERROR: {SRC_MD} not found", file=sys.stderr)
        return 1
    if run_pandoc():
        print(f"Wrote {OUT_DOCX} via pandoc")
        return 0
    print("pandoc not on PATH — falling back to python-docx (lower fidelity).")
    print("For better output install pandoc: https://pandoc.org/installing.html")
    if run_python_docx():
        print(f"Wrote {OUT_DOCX} via python-docx")
        return 0
    print("Failed to build docx", file=sys.stderr)
    return 1


if __name__ == "__main__":
    sys.exit(main())
