"""Regenerate any derived .docx from a source markdown file.

Markdown is the source of truth. The .docx is a derived artifact (gitignored).

Usage:
    python scripts/build_doc.py <source.md> <target.docx>
    # or via Makefile: make docs / make stakeholder-doc / make cpra-letter

Tries pandoc first (better quality output for tables/headings); falls back to
python-docx if pandoc isn't on PATH (less faithful but works without
external dependencies).
"""

from __future__ import annotations

import os
import shutil
import subprocess
import sys


def run_pandoc(src_md: str, out_docx: str) -> bool:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False
    cmd = [pandoc, src_md, "--from", "gfm", "--to", "docx", "-o", out_docx]
    print(f"$ {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"pandoc failed (exit {result.returncode}):\n{result.stderr}",
              file=sys.stderr)
        return False
    return True


def run_python_docx(src_md: str, out_docx: str) -> bool:
    try:
        import docx  # noqa: F401
    except ImportError:
        print("Installing python-docx (one-time)...")
        subprocess.check_call([sys.executable, "-m", "pip", "install",
                               "--quiet", "python-docx"])
    from docx import Document

    doc = Document()
    with open(src_md, encoding="utf-8") as f:
        text = f.read()

    # Crude markdown-to-docx (paragraph + heading level only). Use pandoc
    # for tables/lists fidelity. This fallback is just a backstop.
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == "---":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(out_docx)
    return True


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: build_doc.py <source.md> <target.docx>", file=sys.stderr)
        return 2
    src_md, out_docx = sys.argv[1], sys.argv[2]
    if not os.path.exists(src_md):
        print(f"ERROR: {src_md} not found", file=sys.stderr)
        return 1
    if run_pandoc(src_md, out_docx):
        print(f"Wrote {out_docx} via pandoc")
        return 0
    print("pandoc not on PATH — falling back to python-docx (lower fidelity).")
    print("For better output install pandoc: https://pandoc.org/installing.html")
    if run_python_docx(src_md, out_docx):
        print(f"Wrote {out_docx} via python-docx")
        return 0
    print("Failed to build docx", file=sys.stderr)
    return 1


if __name__ == "__main__":
    sys.exit(main())
